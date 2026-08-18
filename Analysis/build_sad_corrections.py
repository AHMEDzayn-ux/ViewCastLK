"""Build the SAD v2.0 -> v2.1 correction pack.

The SAD was written against SRS v1.1 (212 identifiers). The SRS was then
rewritten as v1.0 with 103 identifiers and fresh numbering, which broke 85 of
the SAD's 158 citations outright and silently repointed most of the rest at
the wrong requirement. This script emits the remap and every consequent edit.

Policy for this pass, per the project owner: THE SRS GOVERNS. Where the two
documents disagree the SAD is changed, not the SRS.
"""
import os
import re

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "Deliverables",
                   "ViewCastLK_SAD_v2.1_Correction_Pack.docx")

# ---------------------------------------------------------------- the remap
# old (SRS v1.1, what the SAD cites)  ->  new (SRS v1.0, what governs now)
# "-" means the requirement has no successor: the citation must be deleted.
REMAP = [
    # --- functional: forecast request ------------------------------------
    ("FR-01", "FR-01", "Planned-video metadata accepted as the basis for a forecast"),
    ("FR-04", "FR-02", "Omitted optional inputs carried through as genuinely missing"),
    ("FR-06", "UR-04", "Invalid input reported by naming the field"),
    ("FR-07", "UR-04", "Entered values preserved across an error or a result"),
    # --- channel statistics ----------------------------------------------
    ("FR-08", "FR-03", "Channel statistics retrieved automatically"),
    ("FR-09", "FR-03", "Channel age derived (folded into FR-03)"),
    ("FR-11", "-", "Error-message wording; SRS §3 now leaves this to design"),
    ("FR-12", "FR-08", "Hidden subscriber count treated as a missing input"),
    # --- title tone -------------------------------------------------------
    ("FR-13", "FR-04", "Four tone dimensions scored"),
    ("FR-16", "-", "Validate/retry behaviour; now SRS §3.9.3 table, not a requirement"),
    ("FR-17", "FR-08", "Forecast returned when tone scores unavailable"),
    ("FR-18", "DB-05", "One tone score set per video"),
    ("FR-19", "FR-05", "No divisiveness score computed, stored, exposed or acted upon"),
    # --- feature assembly -------------------------------------------------
    ("FR-20", "FR-07", "Feature vector assembled"),
    ("FR-21", "FR-07", "Identical transformations at training and serving"),
    ("FR-22", "FR-07", "Frozen encoding map used, not recomputed"),
    ("FR-24", "FR-07", "Published feature order applied"),
    ("FR-25", "-", "No derived feature exposed to the creator; not carried forward"),
    # --- forecasting ------------------------------------------------------
    ("FR-26", "FR-06", "Four horizons predicted from pre-publication features"),
    ("FR-27", "DC-04", "No feature knowable only after publication"),
    ("FR-28", "FR-06", "Single multi-output model, single inference"),
    ("FR-29", "FR-06", "Predictions returned on the real view-count scale"),
    ("FR-30", "FR-08", "Forecast produced with missing inputs, disclosed as such"),
    ("FR-31a", "-", "INTERVALS OUT OF SCOPE - delete the citation and the element"),
    ("FR-31b", "-", "INTERVALS OUT OF SCOPE - delete the citation and the element"),
    # --- presentation -----------------------------------------------------
    ("FR-32", "FR-09", "Trajectory presented graphically and in prose"),
    ("FR-35", "FR-08", "Incomplete-input state indicated on the result view"),
    ("FR-36", "-", "INTERVALS OUT OF SCOPE - FR-09 now forbids interval display"),
    ("FR-36a", "-", "INTERVALS OUT OF SCOPE"),
    ("FR-36b", "-", "INTERVALS OUT OF SCOPE"),
    # --- recommendations --------------------------------------------------
    ("FR-37", "FR-10", "Publishing day and time slot recommended"),
    ("FR-41", "FR-11", "Every recommendation carries its supporting historical pattern"),
    ("FR-42", "FR-11", "Recommendation shown only where ablation shows contribution"),
    ("FR-43", "FR-12", "No guidance advocating divisiveness or misleading claims"),
    # --- accuracy reporting -----------------------------------------------
    ("FR-44", "FR-13", "Accuracy reported combined and per horizon against baseline"),
    ("FR-48", "-", "Outcome tracking REMOVED by SRS §3.1.6 as a deliberate decision"),
    # --- collection -------------------------------------------------------
    ("FR-49", "FR-14", "Verified Sri Lankan channel roster maintained"),
    ("FR-50", "FR-14", "Size tier computed relative to category"),
    ("FR-51", "FR-15", "Discovery via uploads playlists, batched low-cost calls"),
    ("FR-52", "FR-15", "Search-class calls not used for routine discovery"),
    ("FR-53", "DB-01", "Publication timestamp and metadata recorded per video"),
    ("FR-54", "REL-05", "Roster ingestion safely repeatable"),
    ("FR-55", "FR-16", "Engagement polled on a repeating unattended schedule"),
    ("FR-57", "FR-17", "Age computed from the video's own publication timestamp"),
    ("FR-58", "DB-03", "Idempotent writes to the engagement history"),
    ("FR-59", "FR-16", "Video retained until it leaves the tracking window"),
    ("FR-60", "REL-05", "A run survives an individual unavailable video"),
    ("FR-61", "DB-02", "Boundary deviation recorded; target-quality flags"),
    ("FR-62", "FR-18", "All API calls metered through one guard against the allowance"),
    ("FR-63", "FR-15", "Search-class calls permitted only above a safety margin"),
    ("FR-64", "FR-18", "Collection halts cleanly when the allowance is exhausted"),
    ("FR-65", "-", "Retry-with-backoff; now SRS §3.9.3 table, not a requirement"),
    ("FR-66", "SUP-06", "Exactly one run summary per run, including on failure"),
    ("FR-67", "-", "External failure notification; not carried forward"),
    # --- training ---------------------------------------------------------
    ("FR-68", "FR-19", "Derived feature view, one row per video"),
    ("FR-70", "DB-06", "Videos under thirty days retained with horizons absent"),
    ("FR-72", "SUP-05", "Row-filter counts recorded per training run"),
    ("FR-74", "DC-04", "No feature from own or contemporaneous post-publication data"),
    ("FR-75", "FR-20", "Split by publication date; no random cross-validation"),
    ("FR-76", "FR-20", "Target encoding computed inside the training portion only"),
    ("FR-77", "FR-21", "Naive category baseline measured before the model"),
    ("FR-78", "FR-13", "Proportional error reported combined and per horizon"),
    ("FR-80", "FR-22", "Ablation study measures each feature group"),
    ("FR-81", "FR-23", "Reproduction artefacts published with every released model"),
    ("FR-82", "FR-23", "Model versioned and retrained as labels mature"),
    ("FR-82a", "-", "INTERVAL COVERAGE - out of scope"),
    ("FR-82b", "-", "INTERVAL WIDTH - out of scope"),
    ("FR-82c", "-", "INTERVAL RECALIBRATION - out of scope"),
    # --- usability --------------------------------------------------------
    ("UR-01", "UR-01", "Unassisted first use"),
    ("UR-04", "UR-07", "Verified by an informal usability check"),
    ("UR-05", "UR-03", "Everyday language, no statistical notation"),
    ("UR-06", "UR-03", "Numeric accuracy accompanied by a plain explanation"),
    ("UR-07", "FR-11", "Recommendation presented with its supporting pattern"),
    ("UR-08", "UR-04", "Only inputs that cannot be derived automatically"),
    ("UR-09", "UR-04", "Invalid input reported by naming the field"),
    ("UR-10", "UR-04", "Entered values preserved"),
    ("UR-14", "UR-08", "Mobile-sized viewport (Desirable)"),
    # --- reliability ------------------------------------------------------
    ("REL-01", "REL-01", "No guaranteed continuous availability for the dashboard"),
    ("REL-02", "REL-02", "Collection runs to schedule; incomplete runs visible in the log"),
    ("REL-03", "REL-03", "Run success rate and missing-observation rate recorded"),
    ("REL-04", "REL-04", "Scheduler dormancy prevented"),
    ("REL-05", "DB-03", "Idempotent on the natural key"),
    ("REL-06", "REL-05", "Part-way failure leaves the database consistent"),
    ("REL-07", "REL-05", "Warehouse backed up automatically (also DB-07)"),
    ("REL-08", "REL-05", "Unavailable videos handled without failing the run"),
    ("REL-09", "REL-06", "Title-scoring failure degrades the forecast"),
    ("REL-10", "REL-06", "Channel-lookup failure degrades the forecast"),
    ("REL-11", "REL-06", "No placeholder figure when no forecast is possible"),
    ("REL-12", "REL-07", "Quota exhaustion stops collection cleanly"),
    ("REL-13", "REL-08", "No interpolation or smoothing at ingest"),
    ("REL-14", "DB-02", "Boundary deviation recorded; label marked unusable"),
    ("REL-15", "REL-09", "Quantified reduction in error against the naive baseline"),
    ("REL-16", "REL-09", "Measured on held-out data, per horizon and combined"),
    ("REL-18", "REL-11", "Availability and MTBF/MTTR figures to be measured"),
    # --- performance ------------------------------------------------------
    ("PERF-01", "PERF-01", "Immediate acknowledgement and progress indication"),
    ("PERF-02", "PERF-02", "No fixed maximum response time"),
    ("PERF-03", "PERF-03", "Per-call timeout; input treated as missing"),
    ("PERF-05", "-", "Concurrent-volume requirement not carried forward"),
    ("PERF-06", "PERF-04", "Live requests consume negligible quota"),
    ("PERF-07", "PERF-05", "Storage growth monitored against the allocation"),
    ("PERF-08", "PERF-06", "Training feasible on commodity hardware"),
    ("PERF-09", "PERF-06", "Prediction path within serverless limits (merged)"),
    # --- security ---------------------------------------------------------
    ("SEC-01", "SEC-01", "Secrets via the platform secret mechanism"),
    ("SEC-02", "SEC-01", "No secret in source control, client code, logs or responses"),
    ("SEC-03", "SEC-02", "External calls made server-side only"),
    ("SEC-04", "SEC-02", "API keys restricted to the services they serve"),
    ("SEC-06", "SEC-04", "No account, authentication or personal data"),
    ("SEC-07", "SEC-04", "Only publicly available data stored"),
    ("SEC-08", "SEC-05", "Forecast requests rate-limited"),
    ("SEC-10", "SEC-06", "Error messages disclose no implementation detail"),
    # --- supportability ---------------------------------------------------
    ("SUP-01", "SUP-01", "PEP 8 and consistent naming"),
    ("SUP-02", "SUP-01", "Documented repository structure (merged)"),
    ("SUP-03", "SUP-02", "Pinned dependency versions"),
    ("SUP-04", "SUP-03", "Feature preparation implemented once, shared"),
    ("SUP-05", "SUP-04", "Round-trip test on a known training row"),
    ("SUP-06", "SUP-03", "New model inputs added to the shared derived view"),
    ("SUP-07", "SUP-05", "Training run records seed, split boundary, filter counts"),
    ("SUP-08", "FR-23", "Released model versioned with its metrics report"),
    ("SUP-09", "SUP-06", "Run summary sufficient to diagnose a failure"),
    ("SUP-10", "SUP-07", "Configuration and credentials externalised"),
    ("SUP-13", "SUP-08", "Known limitations documented (also DOC-02)"),
    # --- design constraints ----------------------------------------------
    ("DC-01", "DC-01", "Free-tier operation"),
    ("DC-02", "DC-01", "Exceeding a limit degrades rather than charges (merged)"),
    ("DC-06", "DC-03", "Gradient-boosted decision trees"),
    ("DC-07", "DC-04", "No post-publication information in any feature"),
    ("DC-08", "DC-04", "No same-period peer ranking (merged)"),
    ("DC-09", "DC-04", "Backward-looking aggregates end strictly before publication"),
    ("DC-10", "DC-05", "Age computed from the stored publication timestamp"),
    ("DC-11", "DC-05", "Publication-anchored age is the key (merged)"),
    ("DC-12", "DC-06", "UTC storage; local conversion at the edges"),
    ("DC-13", "DC-07", "No divisiveness feature anywhere in the system"),
    ("DC-14", "DC-07", "Guidance only from sustainable levers (merged)"),
    # --- interfaces -------------------------------------------------------
    ("IF-01", "IF-01", "Four dashboard views"),
    ("IF-06", "IF-05", "Current desktop or mobile browser with JavaScript"),
    ("IF-10", "IF-08", "Model layer consumes only the derived feature view"),
    ("IF-13", "IF-10", "Transaction-mode pooler for scheduled jobs"),
    ("IF-14", "IF-11", "Pipeline exposes no inbound network interface"),
    # --- database ---------------------------------------------------------
    ("DB-01", "DB-01", "Relational database as the single authoritative store"),
    ("DB-02", "DB-01", "Holds roster, videos, engagement, tone, quota, run summaries"),
    ("DB-03", "DB-02", "Derived view of one row per video"),
    ("DB-04", "DB-03", "One authoritative observation per publication-anchored day"),
    ("DB-05", "DB-04", "Publication timestamp the source of truth for age"),
    ("DB-06", "DB-05", "Tone scores stored once per video"),
    ("DB-07", "DB-04", "All timestamps stored in UTC (merged)"),
    ("DB-08", "DB-06", "Partial horizon coverage is a normal state"),
    ("DB-09", "DB-07", "Automatic backup; storage growth monitored"),
    ("DB-10", "DB-08", "Transaction-mode pooler for short-lived jobs"),
    ("DB-11", "DB-09", "Observation source recorded"),
    ("DB-12", "DB-10", "Retention policy; unavailable videos marked"),
    # --- legal / docs / standards ----------------------------------------
    ("LEG-01", "LEG-01", "YouTube API terms and developer policies"),
    ("LEG-02", "LEG-02", "No automated retrieval where prohibited; OSS licences"),
    ("LEG-06", "LEG-04", "No personal data collected from dashboard users"),
    ("LEG-07", "LEG-05", "Dataset not redistributed beyond permitted terms"),
    ("DOC-01", "DOC-01", "In-context explanation on each input field"),
    ("DOC-05", "DOC-03", "Help reachable without losing entered values"),
    ("STD-01", "STD-01", "IEEE Std 830-1998 organisation"),
    ("STD-04", "STD-02", "UTF-8, TLS 1.2+, UTC"),
    ("STD-05", "STD-03", "WCAG 2.1 AA (Desirable)"),
]

# ------------------------------------------------------- section-level edits
EDITS = [
    ("§1.4 References",
     'All sixteen figures ... held in the project repository under docs/architecture/puml/',
     "FACTUAL ERROR. That directory does not exist in the repository. Either commit "
     "the sixteen .puml sources to ViewCastLK/docs/architecture/puml/ so the sentence "
     "becomes true, or delete the sentence. Committing them is strongly preferred: it "
     "is the only thing that makes the figures reproducible by anyone but their author.",
     "blocking"),
    ("§1.4 References",
     '"rendered in monochrome with unfilled shapes on a white background, as required '
     'by the course template"',
     "FACTUAL ERROR. Every class box, use-case ellipse and table header carries "
     "PlantUML's default #F1F1F1 grey fill. The figures are monochrome but they are not "
     "unfilled. Apply the skinparam block supplied with this pack and regenerate, or "
     "soften the sentence to 'rendered in monochrome on a white background'.",
     "blocking"),
    ("§2.4 Architectural style",
     '"required by SRS REL-09 through REL-12 and by the degradation-modes table in SRS §3.4.3"',
     "Replace with: 'required by SRS REL-06 and REL-07 and by the failure behaviour "
     "recorded per interface in SRS §3.9.3'. The new SRS has no §3.4.3 and no "
     "degradation-modes table; that content now lives in the Software Interfaces table.",
     "blocking"),
    ("§2.6 Traceability convention",
     '"the composition of the snapshot key is annotated (DB-04, DC-11)"',
     "Replace with (DB-03, DC-05).",
     "blocking"),
    ("§3.1 Goals, row 'Genuine pre-publication forecasting'",
     "Source column: FR-26, FR-27, DC-07, DC-08, DC-09",
     "Replace with: FR-06, DC-04.",
     "blocking"),
    ("§3.1 Goals, row 'No training-serving divergence'",
     "Source column: FR-21, FR-22, FR-24, SUP-04, SUP-05",
     "Replace with: FR-07, SUP-03, SUP-04.",
     "blocking"),
    ("§3.1 Goals, row 'A forecast survives partial failure'",
     "Source column: FR-30, REL-09, REL-10, PERF-03",
     "Replace with: FR-08, REL-06, PERF-03.",
     "blocking"),
    ("§3.1 Goals, row 'Verified Sri Lankan collection'",
     "Source column: FR-49, FR-51, FR-52",
     "Replace with: FR-14, FR-15.",
     "blocking"),
    ("§3.1 Goals, row 'The engagement history is never silently lost'",
     "Source column: FR-58, FR-66, REL-05, REL-06, REL-08",
     "Replace with: DB-03, SUP-06, REL-05.",
     "blocking"),
    ("§3.1 Goals, row 'Labels carry consistent meaning'",
     "Source column: FR-57, FR-61, DC-10, DC-11, REL-14",
     "Replace with: FR-17, DB-02, DB-03, DC-05.",
     "blocking"),
    ("§3.1 Goals, row 'The model technology is not yet fixed'",
     "Source column: DC-06, FR-28",
     "Replace with: DC-03, FR-06. DELETE the sentence 'Interval estimation sits behind "
     "a second interface with two implementations.'",
     "blocking"),
    ("§3.1 Goals, row 'Honest accuracy reporting'",
     "Source column: FR-77 to FR-80, FR-82a to FR-82c",
     "Replace with: FR-13, FR-21, FR-22. DELETE 'empirical interval coverage, interval "
     "width' from the description and DELETE the sentence 'A model with miscalibrated "
     "intervals cannot be released.'",
     "blocking"),
    ("§3.2 Constraints, row 'Fixed daily API quota'",
     "Source column: SRS §2.4, FR-62, FR-63",
     "Replace with: SRS §2.4, FR-15, FR-18.",
     "blocking"),
    ("§3.2 Constraints, row 'Rate-limited language-model service'",
     "Source column: FR-16, SRS §2.4",
     "Replace with: SRS §2.4, §3.9.3. (Old FR-16 has no successor requirement; the "
     "retry-once behaviour is now recorded in the Software Interfaces table.)",
     "blocking"),
    ("§3.2 Constraints, row 'Pre-publication information only'",
     "Source column: DC-07, DC-08",
     "Replace with: DC-04, SRS §3.6.",
     "blocking"),
    ("§3.2 Constraints, row 'Dataset of thousands, not millions'",
     "Source column: SRS §2.4, DC-06",
     "Replace with: SRS §2.4, DC-03.",
     "blocking"),
    ("§3.2 Constraints, row 'Constrained serverless runtime'",
     "Source column: PERF-09, SRS §3.6",
     "Replace with: PERF-06, SRS §3.6.",
     "blocking"),
    ("§3.2 Constraints, row 'Bounded database allocation'",
     "Source column: PERF-07, FR-59, DB-09",
     "Replace with: PERF-05, FR-16, DB-07.",
     "blocking"),
    ("§3.2 Constraints, row 'Ethical constraint on guidance'",
     "Source column: DC-13, DC-14, FR-19, FR-43",
     "Replace with: DC-07, FR-05, FR-12.",
     "blocking"),
    ("§3.2 Constraints, row 'Third-party terms of service'",
     "Source column: LEG-01, LEG-02, DB-12",
     "Replace with: LEG-01, LEG-02, DB-10.",
     "blocking"),
    ("§4.3.1 UC-01, main flow step 2",
     '"reports any field-level problem in plain language (FR-06, UR-09)"',
     "Replace with (UR-04). Both old identifiers now mean something else entirely: "
     "new FR-06 is the four-horizon prediction requirement.",
     "blocking"),
    ("§4.3.1 UC-01, main flow step 3",
     '"retrieves subscriber count ... deriving channel age (FR-08, FR-09)"',
     "Replace with (FR-03).",
     "blocking"),
    ("§4.3.1 UC-01, main flow step 4",
     '"submits the title for tone scoring and validates the returned score set (FR-13 to FR-16)"',
     "Replace with (FR-04). New FR-13 to FR-16 are accuracy reporting, roster "
     "maintenance, discovery and polling - nothing to do with tone.",
     "blocking"),
    ("§4.3.1 UC-01, main flow step 7",
     '"The interval estimator derives a lower and an upper bound for each horizon from '
     'the published per-horizon calibration (FR-31a, FR-31b)."',
     "DELETE this step entirely and renumber the steps that follow. SRS §1.2 places "
     "prediction intervals out of scope and FR-09 requires the trajectory to be "
     "presented without them.",
     "intervals"),
    ("§4.3.1 UC-01, 'Requirements realised' row",
     "FR-01 to FR-06, FR-08 to FR-11, FR-13 to FR-17, FR-20 to FR-31b, FR-32 to FR-36b, "
     "REL-09 to REL-11, PERF-01 to PERF-03",
     "Replace with: FR-01 to FR-04, FR-06 to FR-09, UR-03 to UR-05, REL-06, PERF-01 to PERF-03.",
     "blocking"),
    ("§5.2 Logical view, forecast_domain package",
     "HorizonForecast attributes lower : long [0..1] and upper : long [0..1]",
     "DELETE both attributes from the class and from Figure 5. Keep horizonDays and "
     "point. If the team expects to reinstate intervals later, record that in Appendix A "
     "rather than leaving unused attributes in the design.",
     "intervals"),
    ("§5.2.5 Model inference package",
     "IntervalEstimator interface, QuantileRegressionEstimator and "
     "ResidualQuantileEstimator implementations",
     "DELETE all three classes from the package description and from Figure 5. Also "
     "delete 'the interval calibration' from the frozen-artefact bundle in §2.5 and from "
     "the ModelArtefact class (attribute 'calibration : Map').",
     "intervals"),
    ("§9 Data view, opening paragraph",
     '"SRS DB-02 requires the warehouse to hold title tone scores, a quota ledger and a '
     'run-summary record ... and DB-03 requires a derived view"',
     "Replace the two identifiers with DB-01 and DB-02 respectively.",
     "blocking"),
    ("§9 Data view, video_snapshots key paragraph",
     '"at most one authoritative observation per publication-anchored day per video (DB-04)"',
     "Replace with (DB-03).",
     "blocking"),
    ("§9.1 Stores, run_log row",
     '"Written from a top-level handler so that a crashed run still produces a record '
     '(FR-66, SUP-09)"',
     "Replace with (SUP-06). Status column: 'Required by DB-02' becomes "
     "'Required by DB-01'.",
     "blocking"),
    ("§9.1 Stores, v_video_features row",
     '"(IF-10) ... (SUP-06) ... (DB-08)"',
     "Replace with (IF-08) ... (SUP-03) ... (DB-06). Status column: "
     "'Required by DB-03' becomes 'Required by DB-02'.",
     "blocking"),
    ("§9.2 Invariants, bullet 1",
     "(FR-57, DC-10, DC-11)",
     "Replace with (FR-17, DC-05).",
     "blocking"),
    ("§9.2 Invariants, bullet 2",
     "(FR-58, REL-05)",
     "Replace with (DB-03, REL-05).",
     "blocking"),
    ("§9.2 Invariants, bullet 3",
     "(DC-12, DB-07, STD-04)",
     "Replace with (DC-06, DB-04, STD-02).",
     "blocking"),
    ("§9.2 Invariants, bullet 4",
     "(FR-70, DB-08)",
     "Replace with (DB-06).",
     "blocking"),
    ("§9.2 Invariants, bullet 5",
     "(DB-12, LEG-01)",
     "Replace with (DB-10, LEG-01).",
     "blocking"),
    ("§9.2 Invariants, bullet 6",
     "(DB-11)",
     "Replace with (DB-09).",
     "blocking"),
    ("§9.3 Backup",
     "(REL-07)",
     "Replace with (REL-05, DB-07).",
     "blocking"),
    ("§10 Size and Performance",
     "Whole table plus its opening sentence",
     "Replace with the table supplied in Part 4 of this pack. The current figures are "
     "three days stale and the quota row is presented as measured when it is calculated. "
     "Also update the FR-82 citation on the day-7 label row to FR-23, and the FR-61 "
     "citation on the collection-resolution row to DB-02.",
     "blocking"),
    ("§10, closing paragraph",
     "(SEC-08) ... (PERF-06)",
     "Replace with (SEC-05) ... (PERF-04).",
     "blocking"),
    ("§11 Quality, Reliability row",
     "(REL-03)",
     "No change - REL-03 means the same thing in both versions. Verified.",
     "ok"),
    ("§11 Quality, Recoverability row",
     "(REL-07)",
     "Replace with (REL-05, DB-07).",
     "blocking"),
    ("§11 Quality, Availability row",
     "(REL-01, REL-02)",
     "No change - both retain their meaning. Verified.",
     "ok"),
    ("§11 Quality, Graceful degradation row",
     "(REL-09 to REL-11)",
     "Replace with (REL-06).",
     "blocking"),
    ("§11 Quality, Correctness of labels row",
     "(REL-14, DC-11)",
     "Replace with (DB-02, DC-05).",
     "blocking"),
    ("§11 Quality, Maintainability row",
     "(SUP-01 to SUP-03, SUP-10)",
     "Replace with (SUP-01 to SUP-03, SUP-07).",
     "blocking"),
    ("§11 Quality, Extensibility row",
     '"The model technology sits behind an adapter and interval estimation behind an '
     'interface with two implementations ... (DC-06)"',
     "Replace with: 'The model technology sits behind an adapter, so the open modelling "
     "decision can be resolved without touching the dashboard contract or the feature "
     "layer.' Citation becomes (DC-03).",
     "intervals"),
    ("§11 Quality, Testability row",
     "(SUP-07)",
     "Replace with (SUP-05).",
     "blocking"),
    ("§11 Quality, Observability row",
     "(SUP-09, FR-66)",
     "Replace with (SUP-06).",
     "blocking"),
    ("§11 Quality, Usability row",
     "(UR-01, UR-04)",
     "Replace with (UR-01, UR-07).",
     "blocking"),
    ("§12 References, entry [1]",
     '"ViewCastLK - Software Requirements Specification, Version 1.1"',
     "Replace with 'Version 1.0, 30 July 2026'. The document being tracked is the "
     "rewritten SRS, which is numbered 1.0.",
     "blocking"),
    ("§12 References, entries [21] and [22]",
     "Koenker & Bassett, Regression Quantiles; Lei et al., Distribution-Free Predictive "
     "Inference",
     "DELETE both. Nothing cites them once the interval material is removed. Renumber "
     "[23] to [26] accordingly.",
     "intervals"),
]


# ------------------------------------------------------------------ helpers
def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def body(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(8.5)
            if str(v).strip() == "-" or "out of scope" in str(v).lower() \
                    or str(v).startswith("INTERVAL"):
                run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    return t


# -------------------------------------------------------------------- build
doc = docx.Document()
for s in doc.styles:
    try:
        s.font.name = "Calibri"
    except Exception:
        pass

h(doc, "ViewCastLK — Software Architecture Document", 0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Correction pack: v2.0 → v2.1")
r.bold = True
r.font.size = Pt(14)
body(doc, "Prepared 30 July 2026 · AHAMED M.J.S. (230023E)", size=9)

body(doc,
     "The Software Architecture Document v2.0 was written against SRS version 1.1, which "
     "carried 212 requirement identifiers. The SRS has since been rewritten as version 1.0 "
     "with 103 identifiers and entirely fresh numbering. The architecture document is sound; "
     "its citations are not. Of the 158 requirement identifiers it quotes, 85 no longer exist "
     "and most of the remaining 73 now resolve to a different requirement than the one "
     "intended — which is the more damaging failure, because a reader following the citation "
     "lands somewhere plausible and wrong.")
body(doc,
     "This pack lists every change needed. The governing decision for this revision is that "
     "the SRS wins: where the two documents disagree, the architecture document is amended. "
     "The most significant consequence is that prediction intervals leave the architecture, "
     "because SRS §1.2 places them out of scope and FR-09 requires the trajectory to be "
     "presented without them.")
body(doc,
     "Changes are grouped in four parts. Part 1 is the identifier remap. Part 2 is the "
     "section-by-section edit list. Part 3 covers the removal of the interval subsystem. "
     "Part 4 supplies replacement text for §10, Appendix A and Appendix B.")

doc.add_page_break()

# ---- Part 1
h(doc, "Part 1 — Requirement identifier remap", 1)
body(doc,
     "Every identifier the architecture document quotes, and what it must become. A dash in "
     "the second column means the requirement has no successor in the rewritten SRS: the "
     "citation must be deleted, not repointed. Where several old identifiers collapse onto "
     "one new one, quote the new identifier once rather than repeating it.")
gone = sum(1 for _, n, _ in [(a, b, c) for a, b, c in REMAP] if n == "-")
body(doc, f"{len(REMAP)} identifiers mapped · {len(REMAP) - gone} repointed · "
          f"{gone} deleted outright.", bold=True, size=9)
table(doc, ["SAD v2.0 cites", "Becomes", "Meaning under the rewritten SRS"],
      [(o, n, d) for o, n, d in REMAP])

doc.add_page_break()

# ---- Part 2
h(doc, "Part 2 — Section-by-section edits", 1)
body(doc,
     "Applied in document order. Entries marked ‘intervals’ are consequences of the "
     "scope decision and are collected again in Part 3 so they can be applied as one "
     "coherent change.")
table(doc, ["Where", "Current text", "Change"],
      [(w, c, a) for w, c, a, _ in EDITS])

doc.add_page_break()

# ---- Part 3
h(doc, "Part 3 — Removing the prediction-interval subsystem", 1)
body(doc,
     "SRS §1.2 lists prediction intervals as out of scope, §3.1.5 states that confidence "
     "intervals are not shown in this release, and FR-09 makes that a requirement. The "
     "architecture document currently specifies the subsystem in eleven places. All of them "
     "go together; removing some and not others leaves the document in a worse state than "
     "leaving it alone.")
table(doc, ["Location", "Action"], [
    ("§2.5 Architectural mechanisms, ‘Frozen artefact’ row",
     "Delete ‘the interval calibration’ from the list of what a release publishes."),
    ("§3.1 Goals, ‘The model technology is not yet fixed’",
     "Delete the sentence about interval estimation sitting behind a second interface."),
    ("§3.1 Goals, ‘Honest accuracy reporting’",
     "Delete ‘empirical interval coverage, interval width’ and the sentence about "
     "miscalibrated intervals blocking release."),
    ("§4.3.1 UC-01, main flow",
     "Delete step 7 (the interval estimator) and renumber the following steps."),
    ("§4.3.7 UC-10 Train, calibrate and release",
     "Rename to ‘Train, evaluate and release model’ and remove calibration from the flow "
     "and post-conditions."),
    ("§5.2 forecast_domain, HorizonForecast",
     "Delete attributes lower : long [0..1] and upper : long [0..1]."),
    ("§5.2.5 model_inference",
     "Delete the IntervalEstimator interface and both implementations "
     "(QuantileRegressionEstimator, ResidualQuantileEstimator)."),
    ("§5.2 ModelArtefact",
     "Delete the calibration : Map attribute."),
    ("§6.2 Process group: forecast request",
     "Remove interval derivation from the activity and sequence narratives."),
    ("§11 Quality, Extensibility",
     "Remove the clause about interval estimation behind an interface."),
    ("§12 References",
     "Delete [21] Koenker & Bassett and [22] Lei et al.; renumber what follows."),
    ("Figures 5, 8, 9",
     "Regenerate without the IntervalEstimator classes, the lower/upper attributes and "
     "the interval-derivation step."),
])
body(doc,
     "Record the reversal in Appendix A rather than deleting the history. The replacement "
     "OD-1 in Part 4 does this: the interface is worth reinstating if the SRS scope changes, "
     "and a future reader should be able to see that the option was considered and closed "
     "deliberately rather than never raised.", size=9)

doc.add_page_break()

# ---- Part 4
h(doc, "Part 4 — Replacement text", 1)

h(doc, "4.1  Section 10, Size and Performance", 2)
body(doc,
     "Replace the opening sentence with: ‘The figures below were measured against the "
     "deployed warehouse on 30 July 2026, thirteen days after continuous collection began. "
     "They are the dimensioning characteristics that shaped the architecture, not capacity "
     "targets. One row is calculated rather than measured and is marked as such.’")
table(doc, ["Dimension", "Value", "Basis"], [
    ("Channels tracked", "1,282, all country-verified LK", "Measured"),
    ("Videos captured", "10,392", "Measured"),
    ("Video snapshots", "234,497", "Measured"),
    ("Channel snapshots", "35,460", "Measured"),
    ("Mature day-7 labels", "5,047", "Measured"),
    ("Mature day-14 labels", "156, rising daily", "Measured"),
    ("Storage consumed", "68 MB of a 500 MB allocation, ~5.2 MB per day",
     "Measured"),
    ("Daily quota consumption", "Approximately 8,000 of 10,000 units",
     "CALCULATED from call counts and published unit costs; not yet verified in the "
     "provider console"),
    ("Collection resolution", "Four runs daily at six-hourly intervals", "Configured"),
    ("Batch size", "Fifty videos per statistics call", "Configured"),
])
body(doc,
     "The quota row is the one figure in this table that has never been checked against the "
     "provider's own accounting. Presenting it under a heading that says ‘measured’ invites a "
     "question the team cannot currently answer. Either verify it in the API console before "
     "submission and drop the qualifier, or leave the qualifier in place.", size=9)

h(doc, "4.2  Appendix A, Open Decisions and Conflicts", 2)
body(doc, "Replace the whole appendix with the following five entries.")
table(doc, ["Ref", "Issue", "Position", "Action required"], [
    ("OD-1", "Prediction intervals. SRS §1.2 places them out of scope and FR-09 requires "
             "presentation without them. Earlier drafts of both documents specified lower, "
             "median and upper estimates.",
     "Removed from this document, in accordance with the SRS. The multi-horizon predictor "
     "returns point estimates only.",
     "Reinstate only if the SRS scope changes. Interval coverage and width would then need "
     "to return as release gates, since a range that under-covers is worse than no range."),
    ("OD-2", "Social Blade. The team briefing records it as dropped. SRS §2.1, §3.8, §3.9.3 "
             "and §3.11 retain it as a manual, non-automated import path.",
     "Carried as a manual import path, in accordance with the SRS. No automated interface, "
     "node or scheduled process references it; the observation-source column of the "
     "engagement history (DB-09) distinguishes such imports from automated collection.",
     "If the team confirms it is genuinely dropped, remove the four SRS references and this "
     "entry together."),
    ("OD-3", "Made-for-kids as a required input. SRS FR-01 makes it required. It is true "
             "for 0.27 % of collected videos.",
     "Carried as a required input because FR-01 governs. It is collected and passed to the "
     "model.",
     "If the ablation study required by FR-22 shows no contribution, amend FR-01 to make the "
     "field optional or remove it."),
    ("OD-4", "Prediction runtime. Whether the prediction service is an edge worker "
             "executing an exported artefact or a small service on an equivalent free-tier "
             "platform.",
     "Shown as an edge worker in Figure 12, with the alternative noted. Both satisfy DC-01 "
     "and PERF-06.",
     "Settle once the artefact size and load time are known from a trained model."),
    ("OD-5", "Row-level security on the warehouse. It is currently disabled on all tables.",
     "Accepted for this release. No path reaches the database except server-side code and "
     "the scheduled collector; SEC-02 forbids client-side external calls, and SEC-04 means "
     "no personal data is held.",
     "Revisit if any client-side or third-party read path is ever introduced. Record the "
     "decision in the repository README so it is not mistaken for an oversight."),
])

h(doc, "4.3  Appendix B, Requirements Traceability", 2)
body(doc,
     "Every range in the current matrix overruns the new identifier ceilings — the SRS now "
     "stops at FR-23, UR-08, REL-11, PERF-07, SEC-06, SUP-08, DC-07, DOC-04, IF-11, DB-10, "
     "LEG-05 and STD-03. Replace the matrix with the following.")
table(doc, ["SRS requirement", "Realising architectural element", "Section / Figure"], [
    ("FR-01, FR-02", "ForecastRequest with optional planned day and hour; no imputation "
                     "anywhere in the feature layer", "5.2.1, 5.2.3; Fig. 5"),
    ("FR-03", "ChannelStatsService over YouTubeGateway; failure returned as absent rather "
              "than raised", "5.2.8, 6.2; Fig. 9"),
    ("FR-04, FR-05", "TitleToneService over GeminiGateway; one score row per video; no "
                     "prohibited dimension exists to compute", "5.2.8, 9.1; Fig. 6"),
    ("FR-06 to FR-08", "MultiHorizonPredictor, single inference, inverse transform before "
                       "return; degradation policy in PredictionService",
     "5.2.5, 6.2; Figs. 5, 9"),
    ("FR-09", "ForecastResult composing four HorizonForecast values; presentation layer",
     "5.2.1, 5.2.3"),
    ("FR-10 to FR-12", "RecommendationEngine with AblationGate; no prohibited feature "
                       "exists to optimise against", "5.2.6; Fig. 5"),
    ("FR-13", "MetricsRepository; metrics report published with the artefact", "5.2.7, 4.3.3"),
    ("FR-14, FR-15", "ChannelRosterLoader with country verification; UploadWalker "
                     "incremental traversal", "5.2.9, 4.3.6; Fig. 6"),
    ("FR-16 to FR-18", "CollectionOrchestrator; QuotaGuard ledger and clean halt",
     "5.2.9, 6.3; Figs. 10, 11"),
    ("FR-19 to FR-23", "v_video_features as the tier boundary; time-based splitter; "
                       "baseline and ablation in the evaluation component", "5.2.10, 9.1"),
    ("UR-01 to UR-08", "Dashboard views; stateless request with input retention; UTF-8 "
                       "end to end", "5.2.1, 6.1"),
    ("REL-01 to REL-08", "Idempotent writes; 26-hour look-back; degradation policy confined "
                         "to PredictionService; clean quota halt", "6.2, 6.3, 9.2, 11"),
    ("REL-09 to REL-11", "Evaluation on held-out data gates release; measured operational "
                         "figures recorded for a later revision", "5.2.10, 11"),
    ("PERF-01 to PERF-07", "Immediate acknowledgement; per-call timeouts; batching; artefact "
                           "loaded once per instance", "6.1, 6.2, 7, 10"),
    ("SEC-01 to SEC-06", "Server-side-only external calls; platform secret stores; rate "
                         "limiting; non-disclosing error mapper", "5.2.2, 5.2.8, 7, 11"),
    ("SUP-01 to SUP-08", "Layer rules; shared feature module and round-trip test; versioned "
                         "artefacts; run log; externalised configuration",
     "8.2, 5.2.4, 5.2.10; Fig. 14"),
    ("DC-01 to DC-07", "Free-tier node set; mandated stack; pre-publication feature "
                       "restriction; age invariant in the snapshot key; no prohibited "
                       "feature computed", "3.2, 7, 8.2, 9.2; Figs. 12, 15"),
    ("IF-01 to IF-11", "Four dashboard views; no administrative interface; gateway "
                       "contracts; pooled TLS connections; no inbound pipeline interface",
     "5.2.1, 5.2.8, 7; Figs. 3, 12"),
    ("DB-01 to DB-10", "Eight stores and one derived view; composite snapshot keys; UTC "
                       "storage; partial coverage tolerated; availability marking",
     "9.1, 9.2; Fig. 15"),
    ("LEG-01 to LEG-05", "Retention and attribution enforced in the data layer; no personal "
                         "data store exists", "9.2, 11"),
])
body(doc,
     "Keep the scoping paragraph that follows the matrix, but correct the identifiers it "
     "names: the requirements deliberately excluded as not architecturally significant are "
     "now FR-09 (presentation wording), UR-03 and UR-05, DOC-01 to DOC-04, and STD-01 to "
     "STD-03.", size=9)

doc.add_page_break()

# ---- Part 5
h(doc, "Part 5 — Figures", 1)
body(doc,
     "The sixteen figures are structurally correct and well drawn. Three problems are "
     "presentational and all three are fixable at the source without redrawing anything.")

h(doc, "5.1  Label text is printing at 3.1 to 6.3 point", 2)
body(doc,
     "Every figure is placed at exactly 450 pt (159 mm) wide regardless of its native "
     "proportions. A figure that is 2,046 pixels wide is therefore scaled by 0.22, and "
     "PlantUML's default 14-pixel label text arrives on the page at 3.1 pt — roughly a "
     "quarter of the surrounding body text, and below the size at which print is legible. "
     "Six figures are at 3.1 pt; the best in the document is 6.3 pt.")
table(doc, ["Figures", "Native width", "Label text as printed"], [
    ("5, 6, 7, 11, 13", "2,046 px", "3.1 pt"),
    ("14", "1,963 px", "3.2 pt"),
    ("3, 4, 12", "1,744–1,795 px", "3.5–3.6 pt"),
    ("8, 9, 15", "tall format", "4.3–4.4 pt"),
    ("16", "780 px", "5.0 pt"),
    ("1", "999 px", "6.3 pt"),
])
body(doc,
     "The required source font size is 12 × native_width ÷ 450. For a 2,046 px figure that "
     "is 55 px against a default of 14. Three fixes, in order of preference: raise the "
     "PlantUML font size using the include file supplied with this pack; split the five "
     "figures wider than 1,900 px into two diagrams each; and place the remainder at the "
     "full text-column width rather than 450 pt, which recovers about 30 mm of unused margin.")

h(doc, "5.2  Five figures carry the wrong number inside the image", 2)
body(doc,
     "The figure number is baked into the PlantUML title as well as appearing in the Word "
     "caption. The sections were reordered after the titles were written, so the two now "
     "disagree. The content is correctly placed in every case — only the printed number is "
     "wrong.")
table(doc, ["Page", "Printed inside the image", "Caption says"], [
    ("24", "Figure 7 — Logical decomposition into subsystems", "Figure 5"),
    ("25", "Figure 5 — Architecturally significant classes: serving path", "Figure 6"),
    ("25", "Figure 6 — …classes: collection and training", "Figure 7"),
    ("31", "Figure 10 — Sequence: forecast request", "Figure 9"),
    ("32", "Figure 9 — Activity: one scheduled collection run", "Figure 10"),
])
body(doc,
     "Fix by deleting the title line from all sixteen .puml sources and letting the Word "
     "caption be the only place a figure number exists. Two sources of truth for one number "
     "is what caused this; one source prevents it recurring.")

h(doc, "5.3  The shapes are filled, and §1.4 says they are not", 2)
body(doc,
     "§1.4 states the figures are ‘rendered in monochrome with unfilled shapes on a white "
     "background, as required by the course template’. They are monochrome, but every class "
     "box, use-case ellipse and table header carries PlantUML's default #F1F1F1 grey. Add "
     "the include file below to each source, or soften the sentence.")

h(doc, "5.4  The diagram sources are not in the repository", 2)
body(doc,
     "§1.4 states that the sources are held under docs/architecture/puml/. That directory "
     "does not exist. Until the sixteen .puml files are committed, no one but their author "
     "can regenerate a figure, which defeats the reason for using a text-based tool. This is "
     "the single most important of the four figure issues, because the other three cannot be "
     "fixed by anyone else until it is resolved.")

doc.save(os.path.abspath(OUT))
print("saved", os.path.abspath(OUT))
print(f"remap rows: {len(REMAP)}   section edits: {len(EDITS)}")
