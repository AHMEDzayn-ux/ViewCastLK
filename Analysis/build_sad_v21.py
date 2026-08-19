"""Generate ViewCastLK Software Architecture Document v2.1.

Source of substance: SAD v2.0 (extracted by sad_extract.py).
Governing authority: SRS v1.0 (the rewritten one). Where the two disagree the
architecture document is amended, never the SRS.

Applies:
  * the 158-identifier remap from v1.1 numbering to v1.0 numbering
  * removal of the prediction-interval subsystem (SRS s1.2, FR-09)
  * refreshed s10 figures, measured 30 Jul 2026
  * rewritten Appendix A (five open decisions) and Appendix B (traceability)
  * template compliance: RUP structure, references in s1.4, Letter page,
    'Figure N. caption' form, figures in line with text
"""
import copy
import json
import os
import re

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
BLOCKS = os.path.join(HERE, "sad_blocks.json")
FIGDIR = os.path.join(HERE, "sad_figs")
CROPDIR = os.path.join(HERE, "sad_figs_clean")
TEMPLATE = r"C:\Users\sabit\Downloads\Template for Software Architecture Document.docx"
OUT = os.path.join(HERE, "..", "Deliverables",
                   "ViewCastLK_Software_Architecture_Document_v2.1.docx")

# figure file in page order -> figure number
FIGFILES = ["p08_0", "p10_0", "p11_0", "p16_0", "p24_0", "p25_0", "p25_1",
            "p30_0", "p31_0", "p32_0", "p33_0", "p35_0", "p37_0", "p39_0",
            "p41_0", "p44_0"]
FIG_FOR_FILE = {"p08_0": 1, "p10_0": 2, "p11_0": 3, "p16_0": 4, "p24_0": 5,
                "p25_0": 6, "p25_1": 7, "p30_0": 8, "p31_0": 9, "p32_0": 10,
                "p33_0": 11, "p35_0": 12, "p37_0": 13, "p39_0": 14,
                "p41_0": 15, "p44_0": 16}
FILE_FOR_FIG = {v: k for k, v in FIG_FOR_FILE.items()}

# ---------------------------------------------------------------- remap
REMAP = {
    "FR-01": "FR-01", "FR-04": "FR-02", "FR-06": "UR-04", "FR-07": "UR-04",
    "FR-08": "FR-03", "FR-09": "FR-03", "FR-11": None, "FR-12": "FR-08",
    "FR-13": "FR-04", "FR-16": None, "FR-17": "FR-08", "FR-18": "DB-05",
    "FR-19": "FR-05", "FR-20": "FR-07", "FR-21": "FR-07", "FR-22": "FR-07",
    "FR-24": "FR-07", "FR-25": None, "FR-26": "FR-06", "FR-27": "DC-04",
    "FR-28": "FR-06", "FR-29": "FR-06", "FR-30": "FR-08",
    "FR-31a": None, "FR-31b": None, "FR-32": "FR-09", "FR-35": "FR-08",
    "FR-36": None, "FR-36a": None, "FR-36b": None,
    "FR-37": "FR-10", "FR-41": "FR-11", "FR-42": "FR-11", "FR-43": "FR-12",
    "FR-44": "FR-13", "FR-48": None, "FR-49": "FR-14", "FR-50": "FR-14",
    "FR-51": "FR-15", "FR-52": "FR-15", "FR-53": "DB-01", "FR-54": "REL-05",
    "FR-55": "FR-16", "FR-57": "FR-17", "FR-58": "DB-03", "FR-59": "FR-16",
    "FR-60": "REL-05", "FR-61": "DB-02", "FR-62": "FR-18", "FR-63": "FR-15",
    "FR-64": "FR-18", "FR-65": None, "FR-66": "SUP-06", "FR-67": None,
    "FR-68": "FR-19", "FR-70": "DB-06", "FR-72": "SUP-05", "FR-74": "DC-04",
    "FR-75": "FR-20", "FR-76": "FR-20", "FR-77": "FR-21", "FR-78": "FR-13",
    "FR-80": "FR-22", "FR-81": "FR-23", "FR-82": "FR-23",
    "FR-82a": None, "FR-82b": None, "FR-82c": None,
    "UR-01": "UR-01", "UR-04": "UR-07", "UR-05": "UR-03", "UR-06": "UR-03",
    "UR-07": "FR-11", "UR-08": "UR-04", "UR-09": "UR-04", "UR-10": "UR-04",
    "UR-14": "UR-08",
    "REL-01": "REL-01", "REL-02": "REL-02", "REL-03": "REL-03",
    "REL-04": "REL-04", "REL-05": "DB-03", "REL-06": "REL-05",
    "REL-07": "REL-05", "REL-08": "REL-05", "REL-09": "REL-06",
    "REL-10": "REL-06", "REL-11": "REL-06", "REL-12": "REL-07",
    "REL-13": "REL-08", "REL-14": "DB-02", "REL-15": "REL-09",
    "REL-16": "REL-09", "REL-18": "REL-11",
    "PERF-01": "PERF-01", "PERF-02": "PERF-02", "PERF-03": "PERF-03",
    "PERF-05": None, "PERF-06": "PERF-04", "PERF-07": "PERF-05",
    "PERF-08": "PERF-06", "PERF-09": "PERF-06",
    "SEC-01": "SEC-01", "SEC-02": "SEC-01", "SEC-03": "SEC-02",
    "SEC-04": "SEC-02", "SEC-06": "SEC-04", "SEC-07": "SEC-04",
    "SEC-08": "SEC-05", "SEC-10": "SEC-06",
    "SUP-01": "SUP-01", "SUP-02": "SUP-01", "SUP-03": "SUP-02",
    "SUP-04": "SUP-03", "SUP-05": "SUP-04", "SUP-06": "SUP-03",
    "SUP-07": "SUP-05", "SUP-08": "FR-23", "SUP-09": "SUP-06",
    "SUP-10": "SUP-07", "SUP-13": "SUP-08",
    "DC-01": "DC-01", "DC-02": "DC-01", "DC-06": "DC-03", "DC-07": "DC-04",
    "DC-08": "DC-04", "DC-09": "DC-04", "DC-10": "DC-05", "DC-11": "DC-05",
    "DC-12": "DC-06", "DC-13": "DC-07", "DC-14": "DC-07",
    "IF-01": "IF-01", "IF-06": "IF-05", "IF-10": "IF-08", "IF-13": "IF-10",
    "IF-14": "IF-11",
    "DB-01": "DB-01", "DB-02": "DB-01", "DB-03": "DB-02", "DB-04": "DB-03",
    "DB-05": "DB-04", "DB-06": "DB-05", "DB-07": "DB-04", "DB-08": "DB-06",
    "DB-09": "DB-07", "DB-10": "DB-08", "DB-11": "DB-09", "DB-12": "DB-10",
    "LEG-01": "LEG-01", "LEG-02": "LEG-02", "LEG-06": "LEG-04",
    "LEG-07": "LEG-05",
    "DOC-01": "DOC-01", "DOC-05": "DOC-03",
    "STD-01": "STD-01", "STD-04": "STD-02", "STD-05": "STD-03",
}
IDPAT = re.compile(r"\b(FR|UR|REL|PERF|SEC|SUP|DC|IF|DB|LEG|DOC|STD)-\d+[a-z]?\b")
# whole table rows that exist only to describe the removed interval subsystem
ROW_KILL = re.compile(r"HorizonForecast holds optional bounds|IntervalEstimator",
                      re.I)
RANGEPAT = re.compile(
    r"\b((?:FR|UR|REL|PERF|SEC|SUP|DC|IF|DB|LEG|DOC|STD)-\d+[a-z]?)\s*"
    r"(to|through)\s*((?:FR|UR|REL|PERF|SEC|SUP|DC|IF|DB|LEG|DOC|STD)-\d+[a-z]?)\b")

# explicit whole-cell overrides, applied before token remap
CELL_OVERRIDE = {
    "FR-26, FR-27, DC-07, DC-08, DC-09": "FR-06, DC-04",
    "FR-21, FR-22, FR-24, SUP-04, SUP-05": "FR-07, SUP-03, SUP-04",
    "FR-30, REL-09, REL-10, PERF-03": "FR-08, REL-06, PERF-03",
    "FR-49, FR-51, FR-52": "FR-14, FR-15",
    "FR-58, FR-66, REL-05, REL-06, REL-08": "DB-03, SUP-06, REL-05",
    "FR-57, FR-61, DC-10, DC-11, REL-14": "FR-17, DB-02, DB-03, DC-05",
    "DC-06, FR-28": "DC-03, FR-06",
    "FR-77 to FR-80, FR-82a to FR-82c": "FR-13, FR-21, FR-22",
    "SRS §2.4, FR-62, FR-63": "SRS §2.4, FR-15, FR-18",
    "FR-16, SRS §2.4": "SRS §2.4, §3.9.3",
    "DC-07, DC-08": "DC-04, SRS §3.6",
    "SRS §2.4, DC-06": "SRS §2.4, DC-03",
    "PERF-09, SRS §3.6": "PERF-06, SRS §3.6",
    "PERF-07, FR-59, DB-09": "PERF-05, FR-16, DB-07",
    "DC-13, DC-14, FR-19, FR-43": "DC-07, FR-05, FR-12",
    "LEG-01, LEG-02, DB-12": "LEG-01, LEG-02, DB-10",
    "REL-04, SRS §2.4": "REL-04, SRS §2.4",
    "SRS §2.4, §1.2": "SRS §2.4, §1.2",
}

# ------------------------------------------------- interval / factual surgery
SENTENCE_KILL = [
    r"\s*Interval estimation sits behind a second interface with two implementations\.",
    r"\s*A model with miscalibrated intervals cannot be released\.",
    r"\s*7\. The interval estimator derives a lower and an upper bound for each horizon "
    r"from the published per-horizon calibration \([^)]*\)\.\s*",
    r"\s*6\. Interval bounds are calibrated per horizon and empirical coverage and "
    r"relative width are measured \([^)]*\)\.\s*",
    r"\s*IntervalEstimator is deliberately an interface with two implementations\.",
    r"\s*The SRS specifies bounds derived from the distribution of model errors on "
    r"held-out data, with native quantile regression recorded as the documented "
    r"alternative; the team briefing records quantile regression as decided\.",
    r"\s*The architecture does not resolve that disagreement — it is a modelling "
    r"decision, not a structural one — but it does contain it, so that whichever is "
    r"chosen the serving contract and the dash[^.]*\.",
    r"\s*\[21\] R\. Koenker and G\. Bassett[^\[]*",
    r"\s*\[22\] J\. Lei, M\. G'Sell[^\[]*",
]
TEXT_SUB = [
    # --- interval wording
    (r"feature assembly, inference, interval estimation and recommendation generation",
     "feature assembly, inference and recommendation generation"),
    (r"the encoding map, the ordered feature list, the interval calibration and the "
     r"preprocessing", "the encoding map, the ordered feature list and the preprocessing"),
    (r"computing baseline comparison, per-horizon error, empirical interval coverage, "
     r"interval width and an ablation report",
     "computing baseline comparison, per-horizon error and an ablation report"),
    (r"The model technology sits behind an adapter and interval estimation behind an "
     r"interface with two implementations, so the open model",
     "The model technology sits behind an adapter, so the open model"),
    (r"Inference sits behind an adapter interface\.",
     "Inference sits behind an adapter interface."),
    (r"validation through historical-profile retrieval, feature construction, quantile "
     r"inference and recommendation generation",
     "validation through historical-profile retrieval, feature construction, inference "
     "and recommendation generation"),
    (r"The loop fragment over the four horizons shows that interval estimation is "
     r"per-horizon by construction, because uncertainty grows with dist[^.]*\.",
     "The loop fragment over the four horizons shows that all four values are read from "
     "one multi-output inference rather than obtained from four separate predictors."),
    # --- remaining interval / calibration surgery
    (r"FeatureAssembler, MultiHorizonPredictor and IntervalEstimator all read the same "
     r"ModelArtefact, which is what guarantees that the encoding, the feature order and "
     r"the calibration used at serving time came from the same release as the model "
     r"itself\.",
     "FeatureAssembler and MultiHorizonPredictor read the same ModelArtefact, which is "
     "what guarantees that the encoding and the feature order used at serving time came "
     "from the same release as the model itself."),
    (r"Loads a compatible artefact and produces the four horizon predictions and their "
     r"bounds\. Significant classes: ModelLoader, ModelArtefact, MultiHorizonPredictor, "
     r"IntervalEstimator with its two implementations, ModelVersionMetadata\.",
     "Loads a compatible artefact and produces the four horizon predictions. Significant "
     "classes: ModelLoader, ModelArtefact, MultiHorizonPredictor, ModelVersionMetadata."),
    (r"the ordered feature list, the interval calibration parameters, the shared "
     r"preprocessing implementation", "the ordered feature list, the shared preprocessing "
     "implementation"),
    (r"Train, calibrate and release model", "Train, evaluate and release model"),
    (r"model, encoding map, ordered feature list, calibration, shared preprocessing "
     r"implementation", "model, encoding map, ordered feature list, shared preprocessing "
     "implementation"),
    (r"\s*A model whose measured interval coverage departs materially from nominal is "
     r"recalibrated and is not released until the two agree[^.]*\.", ""),
    (r"Forecast request and result objects, horizon and interval types, recommendation "
     r"types", "Forecast request and result objects, horizon types, recommendation types"),
    (r"Model loading, artefact representation, multi-horizon prediction, interval "
     r"estimation, version metadata",
     "Model loading, artefact representation, multi-horizon prediction, version metadata"),
    (r"a frozen encoding map and published feature order; interval calibration measured "
     r"before release and a release blocked when coverage departs from nominal",
     "a frozen encoding map and published feature order; an automated round-trip test on "
     "a known training row, run before release"),
    (r"Prepares data, fits, evaluates, calibrates and publishes\. Significant classes: "
     r"TrainingPipeline, BaselineModel, Evaluator, IntervalCalibrator, AblationRunner, "
     r"ArtefactPublisher\.",
     "Prepares data, fits, evaluates and publishes. Significant classes: "
     "TrainingPipeline, BaselineModel, Evaluator, AblationRunner, ArtefactPublisher."),
    (r"because SRS\s*(?:FR-\d+[a-z]?)?\s*makes a release conditional on its output: a "
     r"model whose measured interval coverage departs materially from nominal cannot be "
     r"published\.",
     "because the specification makes a release conditional on its output: a model that "
     "does not measurably improve on the naive baseline across all four horizons cannot "
     "be published (REL-09)."),
    (r"A four-horizon forecast with bounds, recommendations and supporting evidence is "
     r"displayed\.",
     "A four-horizon forecast, recommendations and supporting evidence are displayed."),
    (r"A model whose measured interval coverage departs materially from nominal is "
     r"recalibrated and is not released until the two agree[^.]*\.",
     "A model that does not improve on the naive baseline at every horizon is not "
     "released."),
    (r"the trajectory with its band, the plain-language prose",
     "the trajectory, the plain-language prose"),
    (r"UC-10 Train, calibrate and release model", "UC-10 Train, evaluate and release model"),
    (r"4\.3\.7 UC-10 Train, calibrate and release",
     "4.3.7 UC-10 Train, evaluate and release"),
    # --- SRS section pointers that no longer exist
    (r"required by SRS REL-09 through REL-12 and by the degradation-modes table in "
     r"SRS §3\.4\.3", "required by SRS REL-06 and REL-07 and by the failure behaviour "
     "recorded per interface in SRS §3.9.3"),
    (r"Software Requirements Specification, Version 1\.1",
     "Software Requirements Specification, Version 1.0"),
    (r"SRS v1\.1 still references it", "SRS v1.0 still references it"),
    # --- s1.4 factual corrections
    (r"Diagrams are rendered in monochrome with unfilled shapes on a white background, "
     r"as required by the course template\.",
     "Diagrams are rendered in monochrome on a white background, as required by the "
     "course template. A shared style file, docs/architecture/puml/_style.puml, fixes "
     "the label size and removes the default shape fill; figures carried over from "
     "version 2.0 that have not yet been regenerated against it are identified in "
     "Appendix A."),
    # --- stale figures in prose
    (r"as at 27 July 2026, ten days after continuous collection began",
     "as at 30 July 2026, thirteen days after continuous collection began"),
    (r"The figures below are measured observations from the deployed collection pipeline",
     "The figures below were measured against the deployed warehouse"),
]

# ------------------------------------------------------------ replacement tables
SIZE_TABLE = [
    ["Dimension", "Value", "Basis", "Architectural implication"],
    ["Channels tracked", "1,282, all country-verified LK", "Measured",
     "Roster iteration is the outer loop of every run and must remain quota-aware. "
     "Growth in the roster is the primary driver of quota consumption."],
    ["Videos captured", "10,392", "Measured",
     "Video identity writes are idempotent because the same video is rediscovered on "
     "every traversal until it leaves the look-back window."],
    ["Video snapshots", "234,497", "Measured",
     "This table dominates storage and query cost. It is the reason the index set in "
     "Section 9.2 is specified rather than left to implementation."],
    ["Channel snapshots", "35,460", "Measured",
     "Point-in-time channel features require a snapshot-at-or-before join, which is why "
     "channel state is not stored on the video row."],
    ["Mature day-7 labels", "5,047", "Measured",
     "Short-horizon modelling is viable now. This is why training is a repeated "
     "operation rather than a single event (FR-23)."],
    ["Mature day-14 labels", "156, rising daily", "Measured",
     "Day-21 labels begin maturing in early August and day-30 in mid-August. Partial "
     "horizon coverage is therefore the normal state of the warehouse (DB-06)."],
    ["Storage consumed", "68 MB of a 500 MB allocation, growing at about 5.2 MB per day",
     "Measured",
     "At the observed rate the allocation is adequate for the project timeline, but "
     "growth is monitored and the bounded tracking window is what keeps it linear "
     "rather than compounding."],
    ["Daily quota consumption", "Approximately 8,000 of 10,000 units",
     "Calculated, not yet verified in the provider console",
     "A margin of roughly one fifth. This is the figure that justifies reserving the "
     "hundred-unit search endpoint for occasional roster expansion and running two of "
     "the four daily runs in discovery-only mode."],
    ["Collection resolution", "Four runs daily at six-hourly intervals", "Configured",
     "Determines the worst-case deviation between an observation and an exact horizon "
     "boundary, which is why that deviation is recorded per observation and gates label "
     "usability (DB-02)."],
    ["Batch size", "Fifty videos per statistics call", "Configured",
     "Batching, rather than threading, is the concurrency mechanism in the collection "
     "tier."],
]

APPENDIX_A = [
    ["Ref", "Issue", "Position in this document", "Action required"],
    ["OD-1",
     "Prediction intervals. SRS §1.2 places them out of scope and FR-09 requires the "
     "trajectory to be presented without them. Earlier drafts of both documents "
     "specified lower, median and upper estimates, and disagreed on how they should be "
     "derived.",
     "Removed, in accordance with the Software Requirements Specification. The "
     "multi-horizon predictor returns point estimates only, and no interval interface, "
     "calibration artefact or presentation element remains.",
     "Reinstate only if the specification's scope changes. Empirical coverage and "
     "relative width would then have to return as release gates, because a range that "
     "under-covers is more damaging than no range at all."],
    ["OD-2",
     "Social Blade. The team briefing records it as dropped. SRS §2.1, §3.8, §3.9.3 and "
     "§3.11 retain it as a manual, non-automated import path.",
     "Carried as a manual import path, in accordance with the specification. No "
     "automated interface, node or scheduled process references it. The observation "
     "source recorded against each engagement row (DB-09) distinguishes such an import "
     "from automated collection.",
     "If the team confirms it is genuinely dropped, remove the four specification "
     "references and this entry together."],
    ["OD-3",
     "Made-for-kids as a required input. SRS FR-01 makes it a required field. It is "
     "true of 0.27 per cent of collected videos.",
     "Carried as a required input because FR-01 governs. It is collected and passed to "
     "the model.",
     "If the ablation study required by FR-22 shows no measurable contribution, FR-01 "
     "should be amended to make the field optional or to remove it."],
    ["OD-4",
     "Prediction runtime. Whether the prediction service is deployed as an edge worker "
     "executing an exported artefact, or as a small service on an equivalent free-tier "
     "platform.",
     "Shown as an edge worker in Figure 12, with the alternative recorded. Both satisfy "
     "DC-01 and PERF-06, and the choice does not alter any interface in this document.",
     "Settle once a trained artefact exists and its size and load time are known."],
    ["OD-5",
     "Row-level security on the warehouse. It is currently disabled on every table.",
     "Accepted for this release. No path reaches the database except server-side code "
     "and the scheduled collector; SEC-02 forbids client-side external calls and SEC-04 "
     "means no personal data is held.",
     "Revisit if any client-side or third-party read path is introduced. The decision "
     "is recorded in the repository README so that it is not mistaken for an oversight."],
    ["OD-6",
     "Figure regeneration. Figures carried over from version 2.0 were rendered with the "
     "drawing tool's default label size, which prints below the size the course "
     "template requires.",
     "The shared style file docs/architecture/puml/_style.puml sets the label size and "
     "removes the default shape fill. Figures 5, 6, 7, 11 and 13 additionally exceed "
     "the width at which their labels remain legible on a portrait page and are placed "
     "on landscape pages here.",
     "Regenerate all sixteen figures against the style file, and split Figures 5, 6, 7, "
     "11 and 13 into two diagrams each so that neither half needs a landscape page."],
]

APPENDIX_B = [
    ["SRS requirement", "Realising architectural element", "Section / Figure"],
    ["FR-01, FR-02", "ForecastRequest carrying optional planned day and hour; dashboard "
     "form; no imputation anywhere in the feature layer", "5.2.1, 5.2.3; Figure 6"],
    ["FR-03", "ChannelStatsService over YouTubeGateway; failure returned as an absent "
     "value rather than raised", "5.2.8, 6.2; Figure 9"],
    ["FR-04, FR-05", "TitleToneService over GeminiGateway; one score row per video; no "
     "prohibited dimension exists to compute", "5.2.8, 9.1; Figure 7"],
    ["FR-06 to FR-08", "MultiHorizonPredictor producing all four horizons from one "
     "inference; inverse transform before return; degradation policy expressed solely "
     "in PredictionService", "5.2.5, 6.2; Figures 6, 9"],
    ["FR-09", "ForecastResult composing exactly four HorizonForecast values; "
     "presentation layer", "5.2.1, 5.2.3; Figure 6"],
    ["FR-10 to FR-12", "RecommendationEngine gated on the published ablation report; no "
     "prohibited feature exists to optimise against", "5.2.6; Figure 6"],
    ["FR-13", "MetricsRepository; metrics report published with the artefact",
     "5.2.7, 4.3.3"],
    ["FR-14, FR-15", "ChannelRosterLoader with country verification; UploadWalker "
     "incremental traversal; search-class calls gated by QuotaGuard",
     "5.2.9, 4.3.6; Figure 7"],
    ["FR-16 to FR-18", "CollectionOrchestrator; QuotaGuard ledger by call class and "
     "clean halt", "5.2.9, 6.3; Figures 10, 11"],
    ["FR-19 to FR-23", "v_video_features as the declared tier boundary; time-based "
     "splitter; baseline and ablation in the evaluation component", "5.2.10, 9.1"],
    ["UR-01 to UR-08", "Four dashboard views; stateless request with input retention; "
     "UTF-8 end to end", "5.2.1, 6.1; Figure 6"],
    ["REL-01 to REL-08", "Idempotent writes on a natural key; 26-hour look-back; "
     "degradation policy confined to one class; clean quota halt",
     "6.2, 6.3, 9.2, 11; Figures 10, 11"],
    ["REL-09 to REL-11", "Evaluation on held-out data gates release; measured "
     "operational figures recorded for a later revision", "5.2.10, 11"],
    ["PERF-01 to PERF-07", "Immediate acknowledgement and progress indication; per-call "
     "timeouts; batched statistics calls; artefact loaded once per instance",
     "6.1, 6.2, 7, 10"],
    ["SEC-01 to SEC-06", "Server-side-only external calls; platform secret stores; "
     "request rate limiting; non-disclosing error mapper", "5.2.2, 5.2.8, 7, 11"],
    ["SUP-01 to SUP-08", "Layer dependency rules; shared feature module with a "
     "round-trip test; versioned artefacts; run log; externalised configuration",
     "8.2, 5.2.4, 5.2.10, 11; Figure 14"],
    ["DC-01 to DC-07", "Free-tier node set; mandated stack; pre-publication feature "
     "restriction; publication-anchored age inside the snapshot key; no prohibited "
     "feature computed anywhere", "3.2, 7, 8.2, 9.2; Figures 12, 15"],
    ["IF-01 to IF-11", "Four dashboard views; no administrative interface; gateway "
     "contracts; pooled TLS connections; no inbound pipeline interface",
     "5.2.1, 5.2.8, 7; Figures 3, 12"],
    ["DB-01 to DB-10", "Eight stores and one derived view; composite snapshot keys; UTC "
     "storage; partial horizon coverage tolerated; availability marking",
     "9.1, 9.2; Figure 15"],
    ["LEG-01 to LEG-05", "Retention and attribution enforced in the data-access layer; "
     "no personal-data store exists", "9.2, 11"],
]

APP_B_NOTE = (
    "Requirements that are not architecturally significant are deliberately absent. "
    "These are principally the presentation-wording requirement FR-09 and the usability "
    "wording requirements UR-03 and UR-05, the in-context help requirements DOC-01 to "
    "DOC-04, and the standards-conformance requirements STD-01 to STD-03, none of which "
    "constrains the decomposition, the process model, the deployment topology or the "
    "data model. Their absence here is a scoping statement rather than an omission: "
    "they are verified against the specification and the delivered interface, not "
    "against the architecture.")


# ---------------------------------------------------------------- transforms
def fix_ids(text):
    if text.strip() in CELL_OVERRIDE:
        return CELL_OVERRIDE[text.strip()]

    def rng(m):
        a, word, b = m.group(1), m.group(2), m.group(3)
        na, nb = REMAP.get(a), REMAP.get(b)
        if na and nb and na != nb:
            return f"{na} {word} {nb}"
        return na or nb or ""

    text = RANGEPAT.sub(rng, text)
    text = IDPAT.sub(lambda m: REMAP.get(m.group(0), m.group(0)) or "", text)
    # tidy the debris left by deletions
    text = re.sub(r"\(\s*[,;]?\s*\)", "", text)
    text = re.sub(r",\s*,+", ",", text)
    text = re.sub(r"\(\s*,\s*", "(", text)
    text = re.sub(r"\s*,\s*\)", ")", text)
    text = re.sub(r"\s+([,.;)])", r"\1", text)
    # collapse repeats inside one citation list
    def dedupe(m):
        seen, keep = set(), []
        for tok in re.split(r",\s*", m.group(1)):
            t = tok.strip()
            if t and t not in seen:
                seen.add(t)
                keep.append(t)
        return "(" + ", ".join(keep) + ")"
    text = re.sub(r"\(((?:(?:FR|UR|REL|PERF|SEC|SUP|DC|IF|DB|LEG|DOC|STD)-\d+[a-z]?"
                  r"(?:,\s*)?)+)\)", dedupe, text)
    return re.sub(r"[ ]{2,}", " ", text).strip()


def transform(text):
    # the source PDF carries zero-width spaces inside table cells, which silently
    # defeat every pattern below unless they are stripped first
    text = re.sub(r"[​‌‍﻿­]", "", text)
    for pat in SENTENCE_KILL:
        text = re.sub(pat, " ", text)
    for pat, rep in TEXT_SUB:
        text = re.sub(pat, rep, text)
    text = fix_ids(text)
    return re.sub(r"\s{2,}", " ", text).strip()


# ------------------------------------------------------------------ figures
def prepare_figures():
    """Crop the baked-in 'Figure N — ...' title strip off each PNG.

    Five figures print a number that contradicts their caption, because the
    number lives both in the drawing source and in the Word caption and the
    sections were reordered after the sources were written. Removing it from
    the image leaves the caption as the single source of truth.
    """
    os.makedirs(CROPDIR, exist_ok=True)
    meta = {}
    for name in FIGFILES:
        im = Image.open(os.path.join(FIGDIR, name + ".png")).convert("RGB")
        w, h = im.size
        px = im.convert("L").load()
        # find the first fully white row after the title band
        cut = 0
        for y in range(int(h * 0.02), int(h * 0.16)):
            if all(px[x, y] > 245 for x in range(0, w, 7)):
                run = all(all(px[x, yy] > 245 for x in range(0, w, 7))
                          for yy in range(y, min(y + 12, h)))
                if run:
                    cut = y
                    break
        im = im.crop((0, cut, w, h))
        # trim surrounding whitespace so the drawing fills the placed area
        bbox = im.convert("L").point(lambda v: 0 if v > 245 else 255).getbbox()
        if bbox:
            pad = 8
            im = im.crop((max(0, bbox[0] - pad), max(0, bbox[1] - pad),
                          min(im.width, bbox[2] + pad),
                          min(im.height, bbox[3] + pad)))
        out = os.path.join(CROPDIR, name + ".png")
        im.save(out)
        meta[FIG_FOR_FILE[name]] = (out, im.width, im.height)
    return meta


# ------------------------------------------------------------------- docx
def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(ncols):
            set_cell_text(cells[j], row[j] if j < len(row) else "", bold=(i == 0))
    return t


def add_toc_field(par):
    r = par.add_run()
    for el, attr in (("w:fldChar", {"w:fldCharType": "begin"}),
                     ("w:instrText", None),
                     ("w:fldChar", {"w:fldCharType": "separate"}),
                     ("w:t", None),
                     ("w:fldChar", {"w:fldCharType": "end"})):
        e = OxmlElement(el)
        if attr:
            for k, v in attr.items():
                e.set(qn(k), v)
        if el == "w:instrText":
            e.set(qn("xml:space"), "preserve")
            e.text = r'TOC \o "1-3" \h \z \u'
        if el == "w:t":
            e.text = "Right-click and choose Update Field to build the table of contents."
        r._r.append(e)


def landscape_break(doc, on):
    """Start a new section, landscape when on=True, portrait when False."""
    from docx.enum.section import WD_SECTION
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    if on:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = Inches(11), Inches(8.5)
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(0.8)
    else:
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width, s.page_height = Inches(8.5), Inches(11)
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)
    return s


WIDE = {5, 6, 7, 11, 13}   # need a landscape page even after cropping


SENT_SPLIT = re.compile(r'^(.*?[.:;!?])\s+(["“(]?[A-Z].*)$', re.S)


def repair_page_breaks(blocks):
    """Rejoin paragraphs the page breaks tore in half.

    A body block beginning with a lowercase letter is the tail of an earlier
    paragraph. Where the block also carries new sentences after that tail — the
    usual shape when a figure sat on the boundary — only the tail is moved back
    and the remainder stays put, so nothing jumps in front of its figure. The
    search stops at a heading, because a fragment never belongs to a paragraph
    on the far side of one.
    """
    out, joined = [], 0
    for b in blocks:
        if b[0] == "p" and re.match(r"^[a-z]", b[1]) and not b[1].startswith("docs/"):
            j = len(out) - 1
            while j >= 0 and out[j][0] == "tbl":
                j -= 1
            if j >= 0 and out[j][0] in ("p", "fig", "b"):
                m = SENT_SPLIT.match(b[1])
                tail, rest = (m.group(1), m.group(2)) if (
                    m and len(m.group(1)) < len(b[1]) - 20) else (b[1], None)
                host = out[j]
                if host[0] == "fig":
                    out[j] = ("fig", host[1], host[2].rstrip() + " " + tail)
                else:
                    out[j] = (host[0], host[1].rstrip() + " " + tail)
                if rest:
                    out.append(("p", rest))
                joined += 1
                continue
        out.append(b)
    print(f"rejoined {joined} paragraph(s) split by page breaks")
    return out


def main():
    blocks = repair_page_breaks(json.load(open(BLOCKS, encoding="utf8")))
    figs = prepare_figures()

    doc = docx.Document(TEMPLATE)
    doc._body.clear_content()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(1)

    def para(text, style="Body Text", size=10.5):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(text)
        r.font.size = Pt(size)
        return p

    # ---------------------------------------------------------- title page
    for txt, sz, bold in (("University of Moratuwa", 13, False),
                          ("Department of Computer Science and Engineering", 11, False),
                          ("", 8, False),
                          ("ViewCastLK", 26, True),
                          ("A Data-Driven Tool for Forecasting Viewership of "
                           "Sri Lankan YouTube Content", 12, False),
                          ("", 8, False),
                          ("Software Architecture Document", 16, True),
                          ("Version 2.1", 13, False)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(sz)
        r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    add_table(doc, [
        ["Item", "Detail"],
        ["Group / Project", "Group 2 | Project P07"],
        ["Document identifier", "ViewCastLK-SAD-2.1"],
        ["Date", "30 July 2026"],
        ["Governing document",
         "ViewCastLK Software Requirements Specification, Version 1.0"],
    ])

    doc.add_paragraph()
    h = doc.add_heading("Revision History", level=2)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    add_table(doc, [
        ["Date", "Version", "Description", "Author"],
        ["28/Jul/2026", "1.0",
         "Initial architecture document based on the verified collection pipeline.",
         "AHAMED M.U.A. (230025L)"],
        ["30/Jul/2026", "1.1",
         "Data view expanded; open decisions recorded; quality attributes given "
         "verification means.", "AHAMED M.U.A. (230025L)"],
        ["30/Jul/2026", "2.0",
         "Rewritten against the 4+1 view model with architectural mechanisms, "
         "traceability convention and sixteen figures.", "AHAMED M.U.A. (230025L)"],
        ["30/Jul/2026", "2.1",
         "Realigned to Software Requirements Specification v1.0. All requirement "
         "identifiers remapped to the rewritten numbering; the prediction-interval "
         "subsystem removed as out of scope; Section 10 figures re-measured; "
         "Appendices A and B rewritten; references consolidated into Section 1.4 per "
         "the document template.", "AHAMED M.J.S. (230023E)"],
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(
        "Document precedence. The Software Requirements Specification is the authority "
        "on what the system must do. This document is the authority on how it is "
        "structured. Where the two disagree on a requirement, the specification governs "
        "and this document is amended; the disagreements resolved in this revision are "
        "recorded in Appendix A.")
    r.font.size = Pt(9.5)
    r.italic = True

    doc.add_page_break()
    h = doc.add_heading("Table of Contents", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    # ------------------------------------------------------------- body
    skip_until_h1 = None
    cur_landscape = False
    i = 0
    stats = {"h": 0, "p": 0, "tbl": 0, "fig": 0, "b": 0}
    refs_block = None

    # locate the reference list so it can be moved into 1.4 per the template
    for b in blocks:
        if b[0] == "p" and b[1].startswith("[1] Group 2"):
            refs_block = transform(b[1])
    tail = ""
    for b in blocks:
        if b[0] == "p" and re.match(r"^\[1[0-9]\]|\[20\]", b[1]):
            tail += " " + transform(b[1])
    if refs_block:
        refs_block = (refs_block + tail).strip()

    while i < len(blocks):
        b = blocks[i]
        i += 1
        kind = b[0]

        if kind == "h":
            title = transform(b[2] if len(b) > 2 else b[1])
            lvl = b[1]
            # drop section 12 entirely: references now live in 1.4
            if re.match(r"^12\b", title):
                skip_until_h1 = "12"
                continue
            if skip_until_h1 and not title.startswith("Appendix"):
                continue
            skip_until_h1 = None
            if cur_landscape:
                landscape_break(doc, False)
                cur_landscape = False
            hh = doc.add_heading(title, level=min(lvl, 3))
            for r in hh.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
            stats["h"] += 1

            if title.startswith("1.4"):
                para("All documents and online resources cited in this Software "
                     "Architecture Document are listed below in IEEE style. Every web "
                     "resource carries the date on which it was accessed.")
                para("Diagramming tool. All sixteen figures in this document were "
                     "produced with PlantUML using the Graphviz dot layout engine. The "
                     "diagram source is plain text and is held in the project repository "
                     "under docs/architecture/puml/, together with a shared style file "
                     "that fixes label size and shape fill, so that every figure can be "
                     "regenerated deterministically rather than redrawn by hand. "
                     "Diagrams are rendered in monochrome on a white background, as "
                     "required by the course template.")
                if refs_block:
                    for ref in re.split(r"(?=\[\d+\])", refs_block):
                        ref = ref.strip()
                        if ref:
                            rp = doc.add_paragraph(style="Body Text")
                            rp.paragraph_format.space_after = Pt(3)
                            rp.paragraph_format.left_indent = Inches(0.35)
                            rp.paragraph_format.first_line_indent = Inches(-0.35)
                            rr = rp.add_run(ref)
                            rr.font.size = Pt(9.5)
            continue

        if skip_until_h1:
            continue

        if kind == "p":
            txt = transform(b[1])
            if not txt or txt.startswith("[1] Group 2") or re.match(r"^\[1[0-9]\]|\[20\]", txt):
                continue
            if txt.startswith("Table of Contents"):
                continue
            if txt.startswith("University of Moratuwa Department"):
                continue
            if txt.startswith("Revision History"):
                continue
            if txt.startswith("Document precedence."):
                continue
            if len(txt) < 3:
                continue
            para(txt)
            stats["p"] += 1
            continue

        if kind == "b":
            txt = transform(b[1])
            if txt:
                p = doc.add_paragraph(style="Bullet1")
                p.paragraph_format.space_after = Pt(4)
                p.add_run(txt).font.size = Pt(10.5)
                stats["b"] += 1
            continue

        if kind == "tbl":
            rows = [[transform(c) for c in r] for r in b[1]]
            head = " ".join(rows[0]).lower() if rows else ""
            first = " ".join(rows[1]).lower() if len(rows) > 1 else ""
            if "dimension" in head and "architectural implication" in head:
                rows = SIZE_TABLE
            elif rows and rows[0][0].strip() == "Ref" and "od-1" in first:
                rows = APPENDIX_A
            elif "srs requirement" in head:
                rows = APPENDIX_B
            elif rows and rows[0][0] == "Item" and "Group / Project" in " ".join(rows[1]):
                continue
            elif rows and rows[0][0] == "Date" and "Version" in rows[0]:
                continue
            rows = [r for r in rows if any(c.strip() for c in r)]
            rows = [r for r in rows if not ROW_KILL.search(" ".join(r))]
            if not rows:
                continue
            add_table(doc, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            stats["tbl"] += 1
            if rows is APPENDIX_B:
                para(APP_B_NOTE, size=9.5)
            continue

        if kind == "fig":
            num, cap = b[1], transform(b[2])
            path, w, h = figs[num]
            wide = num in WIDE
            if wide and not cur_landscape:
                landscape_break(doc, True)
                cur_landscape = True
            elif not wide and cur_landscape:
                landscape_break(doc, False)
                cur_landscape = False
            avail_w = Inches(9.0) if wide else Inches(6.5)
            avail_h = Inches(6.0) if wide else Inches(7.6)
            scale = min(avail_w / w, avail_h / h)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Emu(int(w * scale)))
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            cr = cp.add_run(f"Figure {num}. {cap}" if cap else f"Figure {num}.")
            cr.font.size = Pt(9.5)
            cr.bold = True
            stats["fig"] += 1
            continue

    if cur_landscape:
        landscape_break(doc, False)

    os.makedirs(os.path.dirname(os.path.abspath(OUT)), exist_ok=True)
    doc.save(os.path.abspath(OUT))
    print("saved", os.path.abspath(OUT))
    print("emitted", stats)


if __name__ == "__main__":
    main()
