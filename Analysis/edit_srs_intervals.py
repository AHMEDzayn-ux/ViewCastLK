"""Amend the ViewCastLK SRS to bring prediction intervals into scope.

Resolves the contradiction with the Software Architecture Document, which
already specifies lower/median/upper estimates. New requirement IDs are
suffixed (FR-31a, FR-36a ...) rather than renumbered, so existing references
from the SAD, test cases and the Gantt remain valid.
"""
import copy
import os

import docx

SRC = r"C:\Users\sabit\Downloads\ViewCastLK_Software_Requirements_Specification.docx"
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "Deliverables",
                   "ViewCastLK_Software_Requirements_Specification_v1.1.docx")

d = docx.Document(SRC)
P = d.paragraphs


def set_text(p, new):
    if not p.runs:
        p.add_run(new)
        return
    p.runs[0].text = new
    for r in p.runs[1:]:
        r.text = ""


def insert_after(anchor, text):
    """Add a Body Text paragraph directly after `anchor`."""
    new = anchor.insert_paragraph_before(text, anchor.style)
    anchor._p.addnext(new._p)
    return new


def add_rows(table, rows, after_last=True):
    """Clone the last data row (keeps formatting) and fill it."""
    for cells in rows:
        tr = copy.deepcopy(table.rows[-1]._tr)
        table._tbl.append(tr)
        for cell, val in zip(table.rows[-1].cells, cells):
            set_text(cell.paragraphs[0], val)
            for extra in cell.paragraphs[1:]:
                set_text(extra, "")


def insert_row_after(table, idx, cells):
    tr = copy.deepcopy(table.rows[idx]._tr)
    table.rows[idx]._tr.addnext(tr)
    row = table.rows[idx + 1]
    for cell, val in zip(row.cells, cells):
        set_text(cell.paragraphs[0], val)
        for extra in cell.paragraphs[1:]:
            set_text(extra, "")


log = []

# ---------------------------------------------------------------- 1. scope
assert "Prediction intervals or confidence bands" in P[105].text
set_text(P[105], "Publishing, scheduling or uploading content to YouTube on the creator's behalf.")
# P[106] already carries that sentence; remove the duplicate just created
if P[106].text.strip().startswith("Publishing, scheduling or uploading"):
    P[105]._p.getparent().remove(P[105]._p)
    log.append("1.2  removed prediction-intervals exclusion from out-of-scope list")
else:
    log.append("1.2  WARNING: check paragraph 105/106 ordering manually")

# ------------------------------------------------------- 2. definitions row
defs = d.tables[1]
for i, r in enumerate(defs.rows):
    if r.cells[0].text.strip() == "Pre-publication forecast":
        insert_row_after(defs, i, [
            "Prediction interval",
            "A lower and an upper bound presented alongside a forecast, calibrated so "
            "that a stated proportion of actual outcomes falls between them."])
        log.append("1.3  added 'Prediction interval' definition")
        break

# ------------------------------------------------- 3. section 3.1.5 forecasting
P = d.paragraphs
anchor = next(p for p in P if p.text.startswith("The model is trained on a transformed scale"))
insert_after(anchor,
    "Alongside each point forecast the system produces a lower and an upper bound. These do not "
    "come from a second model. They are derived from the distribution of the model's own errors, "
    "measured on held-out data the model never trained on. Because training is performed on a "
    "proportional scale, an error quantile expresses itself as a multiplying factor once "
    "predictions are returned to the real scale — a lower bound at some fraction of the projected "
    "figure and an upper bound at some multiple of it. That is the correct shape for a quantity "
    "whose plausible range widens as it grows, and it preserves the single-inference property "
    "described above. Bounds are calibrated separately for each horizon, because uncertainty grows "
    "with distance from the publication date. Native quantile regression is recorded as the "
    "documented alternative should calibrated bounds prove unreliable in measurement; it is not the "
    "first approach because it would require a separate fitted model for each bound at each horizon.")
log.append("3.1.5  added interval-generation paragraph")

add_rows(d.tables[9], [
    ["FR-31a", "The system shall produce a lower and an upper bound alongside the point "
               "forecast at each of the four horizons.", "Essential"],
    ["FR-31b", "The system shall derive interval bounds from the distribution of model errors "
               "measured on held-out data, calibrated separately for each horizon.", "Essential"],
])
log.append("3.1.5  added FR-31a, FR-31b")

# ------------------------------------------------- 4. section 3.1.6 presentation
P = d.paragraphs
excl = next(p for p in P if p.text.startswith("Confidence intervals are not shown"))
set_text(excl,
    "The forecast is presented as a range, not as a single number. Each horizon carries a lower and "
    "an upper figure alongside the projected value, drawn as a band on the trajectory and stated in "
    "ordinary sentences beneath it. The range is given the same prominence as the projected figure "
    "rather than being relegated to a footnote, because a lone number invites a confidence the "
    "underlying accuracy does not support. The interface explains what the range means in plain "
    "terms — that most videos with similar characteristics finished inside it — and does not use "
    "the words confidence interval, quantile or percentile. Where the point forecast can be "
    "produced but bounds cannot, the projected figure is shown by itself and the absence of a range "
    "is stated rather than left for the creator to notice.")
log.append("3.1.6  replaced exclusion paragraph with interval-presentation requirement")

t10 = d.tables[10]
for r in t10.rows:
    if r.cells[0].text.strip() == "FR-36":
        set_text(r.cells[1].paragraphs[0],
                 "The system shall present a lower and an upper bound alongside the projected value "
                 "at every horizon, both on the trajectory chart and in prose.")
        for extra in r.cells[1].paragraphs[1:]:
            set_text(extra, "")
        break
add_rows(t10, [
    ["FR-36a", "The system shall describe the meaning of the range in plain language and shall not "
               "use the terms confidence interval, quantile or percentile in the creator-facing "
               "interface.", "Essential"],
    ["FR-36b", "The system shall present the projected figure alone, and state that no range is "
               "available, when interval bounds cannot be produced.", "Desirable"],
])
log.append("3.1.6  rewrote FR-36; added FR-36a, FR-36b")

# --------------------------------------------- 5. section 3.1.13 evaluation
P = d.paragraphs
anchor = next(p for p in P if p.text.startswith("Success is defined relative to that baseline"))
insert_after(anchor,
    "Interval quality is measured separately from point accuracy, because the metrics that describe "
    "one are silent about the other. A model may place its projected figure well and still state a "
    "range far too narrow to hold the outcomes it claims to cover, and that is the more damaging of "
    "the two failures: a range that promises to contain most results and does not produces exactly "
    "the false confidence the range exists to prevent. Evaluation therefore reports empirical "
    "coverage — the proportion of held-out videos whose actual view count fell within the stated "
    "range — against the proportion the range claims, measured for each horizon separately. It also "
    "reports the typical width of the range relative to the projected figure, because a range wide "
    "enough to contain any plausible outcome is trivially well covered and tells a creator nothing. "
    "Where measured coverage departs materially from what is claimed, the bounds are recalibrated "
    "and the model is not released until the two agree.")
log.append("3.1.13  added interval-evaluation paragraph")

add_rows(d.tables[17], [
    ["FR-82a", "The system shall report empirical interval coverage for each horizon on held-out "
               "data, against the nominal coverage the interval claims.", "Essential"],
    ["FR-82b", "The system shall report the typical interval width relative to the projected "
               "figure, for each horizon.", "Essential"],
    ["FR-82c", "The system shall recalibrate interval bounds when measured coverage departs "
               "materially from nominal, and shall not release a model whose intervals remain "
               "miscalibrated.", "Essential"],
])
log.append("3.1.13  added FR-82a, FR-82b, FR-82c")

# ------------------------------------------------------------- 6. SUP-13
for r in d.tables[35].rows:
    if r.cells[0].text.strip() == "SUP-13":
        set_text(r.cells[1].paragraphs[0],
                 "Known limitations, including the scoping assumption on viewer country, shall be "
                 "documented in a form visible to users of the system.")
        for extra in r.cells[1].paragraphs[1:]:
            set_text(extra, "")
        log.append("SUP-13  removed 'absence of prediction intervals'")
        break

# --------------------------------------------------- 7. 3.9.1 result view
P = d.paragraphs
ui = next(p for p in P if p.text.startswith("The system presents a single web-based user interface"))
set_text(ui, ui.text.replace(
    "The result view presents the four-horizon trajectory graphically and in prose,",
    "The result view presents the four-horizon trajectory graphically and in prose, with the "
    "lower and upper bounds drawn as a band around the projected line and stated alongside it,"))
log.append("3.9.1  result view now describes the interval band")

# ------------------------------------------------------------ 8. references
P = d.paragraphs
last_ref = next(p for p in P if p.text.strip().startswith("[18]"))
r19 = insert_after(last_ref,
    "[19]  R. Koenker and G. Bassett, Jr., “Regression Quantiles,” Econometrica, vol. 46, no. 1, "
    "pp. 33–50, 1978.")
insert_after(r19,
    "[20]  J. Lei, M. G’Sell, A. Rinaldo, R. J. Tibshirani and L. Wasserman, “Distribution-Free "
    "Predictive Inference for Regression,” Journal of the American Statistical Association, "
    "vol. 113, no. 523, pp. 1094–1111, 2018.")
log.append("References  added [19] Koenker & Bassett and [20] Lei et al.")

# ------------------------------------------------------------ revision row
rev = d.tables[0]
add_rows(rev, [["30/Jul/2026", "1.1",
                "Prediction intervals brought into scope, resolving the contradiction with the "
                "Software Architecture Document. Sections 1.2, 1.3, 3.1.5, 3.1.6, 3.1.13, 3.9.1, "
                "SUP-13 and the reference list amended.",
                "AHAMED M.J.S. (230023E)"]])
log.append("Revision history  added the 1.1 row")

os.makedirs(os.path.dirname(os.path.abspath(OUT)), exist_ok=True)
d.save(OUT)
print("saved", os.path.abspath(OUT), "\n")
for line in log:
    print("  •", line)
